from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TARGET_KEYS = ["知识与技能", "过程与方法", "情感态度与价值观"]
PROCESS_TITLE_KEYWORD = "教学过程设计"
OBJECTIVE_LABELS = ["知识与技能目标：", "过程与方法目标：", "情感与价值目标："]
PROCESS_SLOT_CAPACITIES = [38.0, 38.0]
PROCESS_TITLE_CAPACITY_BONUS = 1.1
PROCESS_FIELD_CAPACITY_BONUS = 0.3
PROCESS_STEP_GAP_COST = 0.2
PROCESS_WRAP_WIDTH = 48
PROCESS_PARAGRAPH_LINE_SPACING = 0.95
PROCESS_FIRST_LINE_INDENT_CHARS = 0


def set_run_font(run, font_name: str = "宋体", font_size: float = 10.5, bold: bool = False):
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_paragraph_format(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing: float | None = None,
    first_line_indent_chars: int = 0,
    keep_together: bool = True,
    keep_with_next: bool = False,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if first_line_indent_chars:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_chars * 10.5)
    else:
        paragraph.paragraph_format.first_line_indent = Pt(0)


def replace_paragraph_text(
    paragraph,
    text: str,
    *,
    font_name: str = "宋体",
    font_size: float = 10.5,
    bold: bool = False,
    center: bool | None = None,
):
    paragraph.clear()
    alignment = paragraph.alignment
    if center is not None:
        alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    configure_paragraph_format(paragraph, alignment=alignment)
    run = paragraph.add_run(str(text).replace("\n", " "))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def replace_cell_text(
    cell,
    text: str,
    *,
    font_name: str = "宋体",
    font_size: float = 10.5,
    bold: bool = False,
    center: bool = False,
):
    cell.text = ""
    lines = str(text or "").split("\n")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        configure_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT,
        )
        run = p.add_run(line)
        set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def parse_markdown(md_path: Path) -> Dict[str, object]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    data: Dict[str, object] = {
        "title": "教案（一）",
        "header": {"term": "", "date": "", "weekday": ""},
        "sections": {},
    }

    current_section = None
    current_step = None
    process_steps: List[Dict[str, object]] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            data["title"] = line[2:].strip()
            continue
        if line.startswith("- 学期："):
            data["header"]["term"] = line.split("：", 1)[1].strip()
            continue
        if line.startswith("- 授课日期："):
            date_part = line.split("：", 1)[1].strip()
            if "星期：" in date_part:
                date_value, weekday_value = date_part.split("星期：", 1)
                data["header"]["date"] = date_value.strip()
                data["header"]["weekday"] = weekday_value.strip()
            else:
                data["header"]["date"] = date_part
            continue
        if line.startswith("- 星期："):
            data["header"]["weekday"] = line.split("：", 1)[1].strip()
            continue

        if line.startswith("## "):
            current_section = line[3:].strip()
            current_step = None
            data["sections"][current_section] = process_steps if current_section == "教学过程" else []
            continue

        if line.startswith("### ") and current_section == "教学过程":
            current_step = {"title": line[4:].strip(), "fields": {}}
            process_steps.append(current_step)
            continue

        if current_section == "基本信息" and line.startswith("- ") and "：" in line:
            key, value = line[2:].split("：", 1)
            data["sections"][current_section].append((key.strip(), value.strip()))
            continue

        if current_section == "教学目标" and "：" in line:
            content = line
            if line[:2] in {"1.", "2.", "3.", "4.", "5."}:
                content = line.split(".", 1)[1].strip()
            if content.startswith("- "):
                content = content[2:].strip()
            key, value = content.split("：", 1)
            data["sections"][current_section].append((key.strip(), value.strip()))
            continue

        if current_section == "教学过程" and current_step and line.startswith("- ") and "：" in line:
            key, value = line[2:].split("：", 1)
            current_step["fields"][key.strip()] = value.strip()
            continue

        if current_section:
            data["sections"][current_section].append(line[2:].strip() if line.startswith("- ") else line)

    return data


def extract_docx_content(docx_path: Path) -> Dict[str, object]:
    document = Document(docx_path)

    title = document.paragraphs[0].text.strip() if len(document.paragraphs) > 0 else "教案（一）"
    header_line = document.paragraphs[1].text.strip() if len(document.paragraphs) > 1 else ""
    table0 = document.tables[0]

    def cell_text(row: int, col: int) -> str:
        return table0.cell(row, col).text.strip()

    process_parts: List[str] = []
    for table in document.tables[1:]:
        for row_index in range(len(table.rows)):
            text = table.cell(row_index, 0).text.strip()
            normalized = text.replace(" ", "")
            if text and PROCESS_TITLE_KEYWORD in normalized:
                continue
            if text:
                process_parts.append(text)

    sections = {
        "基本信息": [
            ("课题", cell_text(0, 1)),
            ("课型", cell_text(5, 1)),
            ("课时", cell_text(5, 3)),
            ("使用教具", cell_text(6, 1)),
        ],
        "教学目标": [("知识与技能", cell_text(1, 1))],
        "教学重点": [line for line in cell_text(2, 1).splitlines() if line.strip()],
        "教学难点": [line for line in cell_text(3, 1).splitlines() if line.strip()],
        "突破教学重、难点措施": [line for line in cell_text(4, 1).splitlines() if line.strip()],
        "布置作业": [line for line in cell_text(7, 1).splitlines() if line.strip()],
        "课后小结": [line for line in cell_text(8, 1).splitlines() if line.strip()],
        "教学过程正文": "\n".join(process_parts).strip(),
    }

    return {
        "title": title,
        "header_line": header_line,
        "sections": sections,
    }


def normalize_title(title: str) -> str:
    title = title.strip() or "教案（一）"
    if title == "教案":
        return "教   案  （一）"
    return title


def normalize_date_text(date_text: str, weekday: str) -> str:
    clean = date_text.replace("年", " ").replace("月", " ").replace("日", " ").replace("-", " ").replace("/", " ")
    parts = [part for part in clean.split() if part]
    year = parts[0] if len(parts) > 0 else ""
    month = parts[1].zfill(2) if len(parts) > 1 else ""
    day = parts[2].zfill(2) if len(parts) > 2 else ""

    date_parts = ["授课日期"]
    if year:
        date_parts.append(f"{year}年")
    if month:
        date_parts.append(f"{month}月")
    if day:
        date_parts.append(f"{day}日")
    if weekday:
        date_parts.append(f"星期{weekday}")
    return " ".join(date_parts).strip()


def populate_inline_label_value_paragraph(
    paragraph,
    label: str,
    value: str,
    *,
    label_font_size: float = 10.5,
    value_font_size: float = 10.5,
    center: bool = False,
    right: bool = False,
):
    paragraph.clear()
    alignment = WD_ALIGN_PARAGRAPH.RIGHT if right else WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    configure_paragraph_format(paragraph, alignment=alignment)

    if label:
        label_run = paragraph.add_run(label)
        set_run_font(label_run, font_name="宋体", font_size=label_font_size, bold=False)

    if value:
        value_run = paragraph.add_run(str(value).replace("\n", " "))
        set_run_font(value_run, font_name="宋体", font_size=value_font_size, bold=False)


def build_target_text(target_map: Dict[str, str]) -> str:
    values = [target_map.get(key, "") for key in TARGET_KEYS]
    if any(values):
        lines = [f"{label}{value}" for label, value in zip(OBJECTIVE_LABELS, values) if value]
        return "\n".join(lines)
    if target_map:
        return "\n".join(target_map.values())
    return ""


def fill_first_table(document: Document, data: Dict[str, object]):
    info_map = dict(data["sections"].get("基本信息", []))
    target_map = dict(data["sections"].get("教学目标", []))
    sections = data["sections"]

    table = document.tables[0]
    lesson_topic = info_map.get("课题", "")
    class_name = info_map.get("授课班级", "")

    if lesson_topic:
        replace_cell_text(table.cell(0, 1), lesson_topic)
        table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if class_name:
        replace_cell_text(table.cell(0, 3), class_name)
        table.cell(0, 3).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    target_text = build_target_text(target_map)
    replace_cell_text(table.cell(1, 1), target_text)

    replace_cell_text(table.cell(2, 1), "\n".join(sections.get("教学重点", [])))
    replace_cell_text(table.cell(3, 1), "\n".join(sections.get("教学难点", [])))
    replace_cell_text(table.cell(4, 1), "\n".join(sections.get("突破教学重、难点措施", [])))
    replace_cell_text(table.cell(5, 1), info_map.get("课型", ""), center=True)
    replace_cell_text(table.cell(5, 3), info_map.get("课时", ""), center=True)
    replace_cell_text(table.cell(6, 1), info_map.get("使用教具", ""))
    replace_cell_text(table.cell(7, 1), "")
    replace_cell_text(table.cell(8, 1), "")


def clone_table_after(document: Document, source_table, insert_after_element):
    new_tbl = deepcopy(source_table._tbl)
    insert_after_element.addnext(new_tbl)
    return document.tables[-1]


def insert_page_break_after(element):
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    element.addnext(paragraph)
    return paragraph


def get_process_tables(document: Document):
    process_tables = []
    for table in document.tables[1:]:
        if not table.rows:
            continue
        header_text = table.cell(0, 0).text.replace(" ", "").replace("\n", "")
        if PROCESS_TITLE_KEYWORD in header_text:
            process_tables.append(table)
    return process_tables


def get_process_content_cells(document) -> List[object]:
    cells = []
    for table in get_process_tables(document):
        for row_index in range(1, len(table.rows)):
            if len(table.rows[row_index].cells) == 0:
                continue
            cells.append(table.cell(row_index, 0))
    return cells


def append_process_table_page(document: Document, template_table):
    page_break = insert_page_break_after(document._body._element[-1])
    return clone_table_after(document, template_table, page_break)


def measure_text_units(text: str, wrap_width: int = PROCESS_WRAP_WIDTH) -> float:
    stripped = text.strip()
    if not stripped:
        return 0.0

    width = 0.0
    for char in stripped:
        if char.isspace():
            width += 0.5
        elif ord(char) < 128:
            width += 1.0
        else:
            width += 2.0
    return max(1.0, (width + wrap_width - 1) // wrap_width)


def split_long_text_to_segments(text: str, wrap_width: int = PROCESS_WRAP_WIDTH, target_units: float = 7.5) -> List[str]:
    stripped = text.strip()
    if not stripped:
        return []

    sentences = [part.strip() for part in stripped.replace("\n", "").split("。") if part.strip()]
    rebuilt = [f"{part}。" for part in sentences]
    if not rebuilt:
        rebuilt = [stripped]

    segments: List[str] = []
    current_parts: List[str] = []
    current_units = 0.0
    for sentence in rebuilt:
        sentence_units = measure_text_units(sentence, wrap_width)
        if current_parts and current_units + sentence_units > target_units:
            segments.append("".join(current_parts).strip())
            current_parts = [sentence]
            current_units = sentence_units
        else:
            current_parts.append(sentence)
            current_units += sentence_units
    if current_parts:
        segments.append("".join(current_parts).strip())

    return segments or [stripped]


def estimate_process_line_usage(lines: List[Dict[str, object]], wrap_width: int = PROCESS_WRAP_WIDTH) -> float:
    total = 0.0
    for line in lines:
        text = str(line.get("text", "")).strip()
        if not text:
            continue
        total += measure_text_units(text, wrap_width)
        if line.get("kind") == "step_title":
            total += PROCESS_TITLE_CAPACITY_BONUS
        elif line.get("kind") == "field_label":
            total += PROCESS_FIELD_CAPACITY_BONUS
    return total


def build_process_blocks(sections: Dict[str, object]) -> List[List[Dict[str, object]]]:
    process_text = sections.get("教学过程正文", "")
    if process_text:
        blocks: List[List[Dict[str, object]]] = []
        for paragraph in [part.strip() for part in process_text.split("\n\n") if part.strip()]:
            block_lines = []
            for index, raw_line in enumerate(paragraph.splitlines()):
                text = raw_line.strip()
                if not text:
                    continue
                kind = "step_title" if index == 0 else "body"
                block_lines.append({"kind": kind, "text": text, "bold": kind == "step_title"})
            if block_lines:
                blocks.append(block_lines)
        return blocks

    blocks = []
    for step in sections.get("教学过程", []):
        if isinstance(step, str):
            text = step.strip()
            if text:
                blocks.append([{"kind": "body", "text": text, "bold": False}])
            continue

        fields = step.get("fields", {})
        block_lines: List[Dict[str, object]] = []
        if step.get("title"):
            block_lines.append({"kind": "step_title", "text": str(step["title"]).strip(), "bold": True})

        for field_name in ["教师活动", "学生活动", "设计意图", "时间"]:
            value = str(fields.get(field_name, "")).strip()
            if not value:
                continue
            block_lines.append({"kind": "field_label", "text": f"{field_name}：", "bold": True})
            for segment in split_long_text_to_segments(value):
                block_lines.append({"kind": "body", "text": segment, "bold": False})

        if block_lines:
            blocks.append(block_lines)

    return blocks


def split_block_for_capacity(block: List[Dict[str, object]], capacity: float) -> List[List[Dict[str, object]]]:
    if not block:
        return []

    pieces: List[List[Dict[str, object]]] = []
    current: List[Dict[str, object]] = []
    current_usage = 0.0

    for index, line in enumerate(block):
        line_usage = estimate_process_line_usage([line])
        projected = current_usage + line_usage + (PROCESS_STEP_GAP_COST if current else 0.0)
        if current and projected > capacity:
            pieces.append(current)
            current = []
            current_usage = 0.0

        if not current and index > 0 and line.get("kind") == "body":
            current.append({"kind": "continuation", "text": "（续）", "bold": True})
            current_usage = estimate_process_line_usage(current)

        if current:
            current_usage += PROCESS_STEP_GAP_COST
        current.append(line)
        current_usage += line_usage

    if current:
        pieces.append(current)

    return pieces


def paginate_process_blocks(blocks: List[List[Dict[str, object]]], slot_capacities: List[float]) -> List[List[Dict[str, object]]]:
    pages: List[List[Dict[str, object]]] = []
    current_page: List[Dict[str, object]] = []
    current_usage = 0.0
    slot_index = 0

    def capacity_for(index: int) -> float:
        if index < len(slot_capacities):
            return slot_capacities[index]
        return PROCESS_SLOT_CAPACITIES[index % len(PROCESS_SLOT_CAPACITIES)]

    for block in blocks:
        pending_blocks = [block]
        while pending_blocks:
            current_block = pending_blocks.pop(0)
            capacity = capacity_for(slot_index)
            block_usage = estimate_process_line_usage(current_block)
            extra_gap = PROCESS_STEP_GAP_COST if current_page else 0.0

            if current_page and current_usage + extra_gap + block_usage <= capacity:
                current_usage += extra_gap
                current_page.extend(current_block)
                continue

            if not current_page and block_usage <= capacity:
                current_page.extend(current_block)
                current_usage = block_usage
                continue

            if not current_page:
                split_blocks = split_block_for_capacity(current_block, capacity)
                head = split_blocks[0]
                tail = split_blocks[1:]
                current_page.extend(head)
                current_usage = estimate_process_line_usage(head)
                pending_blocks = tail + pending_blocks
            else:
                pages.append(current_page)
                current_page = []
                current_usage = 0.0
                slot_index += 1
                pending_blocks.insert(0, current_block)
                continue

            if current_page:
                pages.append(current_page)
                current_page = []
                current_usage = 0.0
                slot_index += 1

    if current_page or not pages:
        pages.append(current_page)

    return pages


def clear_cell(cell):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    configure_paragraph_format(paragraph)
    return paragraph


def write_process_lines_to_cell(cell, lines: List[Dict[str, object]]):
    clear_cell(cell)
    first = True
    for line in lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        kind = line.get("kind")
        is_title = kind in {"step_title", "continuation"}
        configure_paragraph_format(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=PROCESS_PARAGRAPH_LINE_SPACING,
            first_line_indent_chars=0 if is_title or kind == "field_label" else PROCESS_FIRST_LINE_INDENT_CHARS,
            keep_with_next=kind in {"step_title", "field_label"},
        )
        run = paragraph.add_run(str(line.get("text", "")))
        set_run_font(run, font_name="宋体", font_size=10.5, bold=bool(line.get("bold", False)))


def fill_process_content(document: Document, data: Dict[str, object], template_path: Path):
    sections = data["sections"]
    process_tables = get_process_tables(document)
    if not process_tables:
        return

    blocks = build_process_blocks(sections)
    if not blocks:
        content_cells = get_process_content_cells(document)
        for cell in content_cells:
            clear_cell(cell)
        return

    template_table = process_tables[0]
    initial_slot_count = len(process_tables) * len(PROCESS_SLOT_CAPACITIES)
    slot_capacities = [
        PROCESS_SLOT_CAPACITIES[index % len(PROCESS_SLOT_CAPACITIES)]
        for index in range(initial_slot_count)
    ]
    pages = paginate_process_blocks(blocks, slot_capacities)
    required_pages = max(1, (len(pages) + len(PROCESS_SLOT_CAPACITIES) - 1) // len(PROCESS_SLOT_CAPACITIES))

    while len(process_tables) < required_pages:
        process_tables.append(append_process_table_page(document, template_table))

    content_cells = get_process_content_cells(document)
    for index, cell in enumerate(content_cells):
        page_lines = pages[index] if index < len(pages) else []
        write_process_lines_to_cell(cell, page_lines)


def build_doc(data: Dict[str, object], template_path: Path, output_path: Path):
    document = Document(template_path)

    title = normalize_title(str(data.get("title", "教案（一）")))
    header = data.get("header", {})
    date_line = normalize_date_text(header.get("date", ""), header.get("weekday", ""))
    if header.get("term"):
        date_line = f"{header['term']}    {date_line}"

    if len(document.paragraphs) >= 1:
        replace_paragraph_text(document.paragraphs[0], title, font_name="黑体", font_size=22, bold=True, center=True)
    if len(document.paragraphs) >= 2:
        populate_inline_label_value_paragraph(document.paragraphs[1], "", date_line, label_font_size=12, value_font_size=12, right=True)

    fill_first_table(document, data)
    fill_process_content(document, data, template_path)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="将教案 Markdown 按原模板版式转换为 Word 文档")
    parser.add_argument("input", nargs="?", help="输入的 Markdown 文件路径")
    parser.add_argument("-o", "--output", help="输出的 Word 文件路径")
    parser.add_argument("-t", "--template", default="教案模板.docx", help="Word 模板路径")
    parser.add_argument("--from-docx", help="从现有 docx 提取内容后再按模板输出")
    args = parser.parse_args()

    template_path = Path(args.template)
    if not template_path.exists():
        raise FileNotFoundError(f"找不到模板文件: {template_path}")

    if args.from_docx:
        source_docx = Path(args.from_docx)
        if not source_docx.exists():
            raise FileNotFoundError(f"找不到来源 docx: {source_docx}")
        data = extract_docx_content(source_docx)
        output_path = Path(args.output) if args.output else source_docx.with_name(f"{source_docx.stem}_按模板输出.docx")
    else:
        if not args.input:
            raise ValueError("未提供 Markdown 输入文件")
        input_path = Path(args.input)
        if not input_path.exists():
            raise FileNotFoundError(f"找不到输入文件: {input_path}")
        data = parse_markdown(input_path)
        output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")

    build_doc(data, template_path, output_path)
    print(f"已生成: {output_path}")


if __name__ == "__main__":
    main()
